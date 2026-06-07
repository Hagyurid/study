from io import BytesIO
from pathlib import Path
from typing import List, Optional
from html import escape
from urllib.parse import quote
import shutil
import tempfile
import zipfile
import base64
import re

from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse, Response
from fastapi.staticfiles import StaticFiles

from app.config import ACTION_API_KEY, PUBLIC_BASE_URL, UPLOAD_DIR
from app.db import (
    create_project, create_workflow_plan, create_workflow_run, delete_source, delete_calculator_blueprint,
    export_project_note_as_source, final_bundle, get_latest_note_version,
    get_dashboard,
    get_mapping_status, get_next_section, get_source_markdown, get_problem_pack_by_token, get_source,
    get_calculator_blueprint, get_unit_map, get_workflow_options, init_db, list_external_notes,
    list_note_versions, list_sources, list_study_notes, list_unit_maps, list_unmapped_sources,
    list_calculator_blueprints, list_workflow_checkpoints, list_workflow_plans, list_workflow_runs,
    save_calculator_blueprint, save_note_version, save_outline, save_problem_pack,
    save_project_items, save_section, save_source, save_text_source,
    save_transcript_revision, save_unit_map, save_workflow_checkpoint, update_source_markdown,
    search_sources, get_next_workflow_step,
)
from app.models import (
    CalculatorBlueprintSave, GenericItemsSave, NoteVersionSave, OutlineSave,
    ProblemPackSave, ProjectCreate, SearchRequest, SectionSave, TextSourceSave,
    TranscriptRevisionSave, UnitMapSave, WorkflowPlanCreate, WorkflowRunCreate,
    WorkflowCheckpointSave, StudyNoteSave, StudyNoteUpdate, ExamCramSave,
)
from app.modules.file_extract import extract_text
from app.modules.prgm_engine import generate_txt_files, validate_blueprint

ROOT = Path(__file__).resolve().parent.parent
STATIC = ROOT / "static"

app = FastAPI(title="LectureNote Suite", version="2.2.8")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"], allow_credentials=False)
app.mount("/static", StaticFiles(directory=str(STATIC)), name="static")
init_db()

SOURCE_TYPES = [
    ("lecture_slides", "강의 슬라이드"), ("textbook", "교재"), ("transcript", "전사본"),
    ("corrected_transcript", "보정 전사본"), ("past_exam", "시험지/기출"),
    ("exam_trend", "시험 정보/경향"), ("generated_note", "GPT 생성 정리본"),
    ("exam_cram", "시험 직전 정리"),
    ("external_note", "외부 정리본"), ("calculator_manual", "계산기 사용법/구조 해설"), ("note_example", "정리본 예시"),
]


def _is_authorized(authorization: Optional[str] = None, x_action_key: Optional[str] = None, action_key: str = "") -> bool:
    if not ACTION_API_KEY:
        return True
    return authorization == f"Bearer {ACTION_API_KEY}" or x_action_key == ACTION_API_KEY or action_key == ACTION_API_KEY


def require_auth(authorization: Optional[str] = Header(default=None), x_action_key: Optional[str] = Header(default=None)):
    if _is_authorized(authorization, x_action_key):
        return True
    raise HTTPException(status_code=401, detail="Invalid action key")


def _options(selected: str = "lecture_slides") -> str:
    return "\n".join(f'<option value="{v}" {"selected" if v == selected else ""}>{label}</option>' for v, label in SOURCE_TYPES)


EXAM_SCOPE_LABELS = {
    "unknown": "시험범위 확인 안 됨",
    "in_scope": "현재 시험범위 해당",
    "out_of_scope": "현재 시험범위 아님",
    "mixed": "일부만 시험범위 해당",
}
EXAM_USAGE_LABELS = {
    "style_generation": "출제 스타일 분석 후 새 문제 생성",
    "exact_transcription": "해당 시험지 문제를 그대로 전사",
    "both": "그대로 전사 + 스타일 기반 새 문제 생성",
    "exclude": "참고 제외/보관만",
}

def _exam_meta_text(exam_scope_status: str, exam_usage_mode: str, exam_range_note: str) -> str:
    scope = EXAM_SCOPE_LABELS.get(exam_scope_status or "unknown", exam_scope_status or "unknown")
    usage = EXAM_USAGE_LABELS.get(exam_usage_mode or "style_generation", exam_usage_mode or "style_generation")
    note = (exam_range_note or "").strip() or "없음"
    return f"""[시험지 메타]
시험범위 해당 여부: {scope}
문제 사용 방식: {usage}
범위/사용 메모: {note}
""".strip()


def _page(title: str, body: str) -> HTMLResponse:
    return HTMLResponse(f"""<!doctype html><html lang="ko"><head><meta charset="utf-8"/><meta name="viewport" content="width=device-width, initial-scale=1"/><title>{escape(title)}</title><style>
:root{{--bg:#f6f7fb;--card:#fff;--text:#171923;--muted:#667085;--line:#e4e7ec;--accent:#4f46e5;--accent2:#eef2ff;--danger:#d92d20;--ok:#047857}}*{{box-sizing:border-box}}body{{margin:0;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,'Noto Sans KR',Arial,sans-serif;background:linear-gradient(180deg,#f8f9ff 0%,#f6f7fb 42%,#f3f4f8 100%);color:var(--text)}}.wrap{{max-width:1120px;margin:0 auto;padding:32px 20px 56px}}.top{{display:flex;justify-content:space-between;align-items:flex-start;gap:16px;margin-bottom:24px}}h1{{font-size:30px;line-height:1.2;margin:0 0 8px;letter-spacing:-.035em}}h2{{margin:0 0 8px;font-size:20px;letter-spacing:-.02em}}p{{margin:0;color:var(--muted);line-height:1.6}}.nav,.actions{{display:flex;gap:10px;flex-wrap:wrap;margin-top:16px}}a.btn,button,.btn{{border:0;border-radius:12px;padding:10px 14px;background:var(--accent);color:#fff;text-decoration:none;font-weight:750;cursor:pointer;display:inline-flex;align-items:center;justify-content:center;gap:6px}}a.btn.secondary,.btn.secondary{{background:var(--accent2);color:var(--accent)}}button.danger{{background:var(--danger)}}.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));gap:16px}}.card{{background:rgba(255,255,255,.92);border:1px solid var(--line);border-radius:22px;padding:22px;box-shadow:0 12px 34px rgba(16,24,40,.06);backdrop-filter:blur(8px)}}label{{display:block;font-weight:760;margin:16px 0 8px}}input,select{{width:100%;border:1px solid var(--line);border-radius:13px;padding:12px 13px;font-size:15px;background:white;outline:none}}input:focus,select:focus{{border-color:#818cf8;box-shadow:0 0 0 4px rgba(79,70,229,.12)}}input[type=file]{{padding:16px;border-style:dashed;background:#fbfcff}}input[type=checkbox]{{width:auto;accent-color:var(--accent)}}.hint{{font-size:13px;color:var(--muted);margin-top:6px}}.keybox{{display:flex;gap:10px;align-items:center;flex-wrap:wrap;margin-top:8px}}.keybox label{{margin:0;font-weight:600;color:var(--muted);font-size:13px;display:flex;gap:6px;align-items:center}}.keybox input[type=checkbox]{{width:auto}}.key-status{{font-size:13px;color:var(--ok);font-weight:700}}table{{width:100%;border-collapse:collapse;background:white;border-radius:16px;overflow:hidden}}th,td{{border-bottom:1px solid var(--line);padding:12px;text-align:left;vertical-align:top;font-size:14px}}th{{background:#f9fafb;color:#344054;font-size:13px}}code{{background:#f2f4f7;padding:2px 6px;border-radius:6px}}.pill{{display:inline-block;padding:4px 8px;border-radius:999px;background:var(--accent2);color:var(--accent);font-size:12px;font-weight:800}}.muted{{color:var(--muted)}}.kbd{{border:1px solid var(--line);background:#fff;border-radius:8px;padding:2px 7px;font-size:12px;color:#344054}}@media(max-width:640px){{.top{{display:block}}.wrap{{padding:22px 14px}}h1{{font-size:25px}}}}</style></head><body><main class="wrap">{body}</main><script>
(function(){{
  const KEY = "lecturenote_action_key";
  const SUBJECT = "lecturenote_subject";
  function qs(sel, root=document){{ return Array.from(root.querySelectorAll(sel)); }}
  function setStatus(msg){{
    qs("[data-key-status]").forEach(el => el.textContent = msg || "");
  }}
  function fill(){{
    const params = new URLSearchParams(window.location.search);
    const queryKey = params.get("action_key") || "";
    const querySubject = params.get("subject") || "";
    if (queryKey) localStorage.setItem(KEY, queryKey);
    if (querySubject) localStorage.setItem(SUBJECT, querySubject);
    const savedKey = localStorage.getItem(KEY) || "";
    const savedSubject = localStorage.getItem(SUBJECT) || "";
    qs("input[name='action_key']").forEach(input => {{
      if (savedKey && !input.value) input.value = savedKey;
    }});
    qs("input[name='subject']").forEach(input => {{
      if (savedSubject && !input.value) input.value = savedSubject;
    }});
    if (savedKey) setStatus("저장된 액션 키를 자동 입력했습니다.");
  }}
  function remember(){{
    qs("input[name='action_key']").forEach(input => {{
      const v = (input.value || "").trim();
      if (v) localStorage.setItem(KEY, v);
    }});
    qs("input[name='subject']").forEach(input => {{
      const v = (input.value || "").trim();
      if (v) localStorage.setItem(SUBJECT, v);
    }});
  }}
  function clearKey(){{
    localStorage.removeItem(KEY);
    qs("input[name='action_key']").forEach(input => input.value = "");
    setStatus("저장된 액션 키를 삭제했습니다.");
  }}
  window.clearLectureNoteActionKey = clearKey;
  document.addEventListener("DOMContentLoaded", () => {{
    fill();
    qs("form").forEach(form => form.addEventListener("submit", remember));
    qs("input[name='action_key']").forEach(input => input.addEventListener("change", remember));
    qs("input[name='subject']").forEach(input => input.addEventListener("change", remember));
    qs("[data-clear-key]").forEach(btn => btn.addEventListener("click", clearKey));
  }});
}})();
</script></body></html>""")


async def _save_uploaded_file(file: UploadFile, source_type: str, subject: str, title: str = "", exam_scope_status: str = "unknown", exam_usage_mode: str = "style_generation", exam_range_note: str = "") -> dict:
    suffix = Path(file.filename or "source.bin").suffix
    source_dir = UPLOAD_DIR / source_type
    source_dir.mkdir(parents=True, exist_ok=True)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    size = 0
    try:
        with open(tmp.name, "wb") as out:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                out.write(chunk)
        safe_name = (file.filename or "file").replace("/", "_").replace("\\", "_")
        stored = source_dir / f"{Path(tmp.name).name}-{safe_name}"
        shutil.move(tmp.name, stored)
        text, status = extract_text(stored, safe_name, file.content_type or "")
        if source_type == "past_exam":
            meta = _exam_meta_text(exam_scope_status, exam_usage_mode, exam_range_note)
            text = meta + "\n\n" + (text or "")
        else:
            exam_scope_status = "unknown"
            exam_usage_mode = "style_generation"
            exam_range_note = ""
        inferred_title = title.strip() or Path(safe_name).stem or "untitled"
        return save_source(inferred_title, source_type, safe_name, str(stored), file.content_type or "", size, text, status, subject=subject.strip(), exam_scope_status=exam_scope_status, exam_usage_mode=exam_usage_mode, exam_range_note=exam_range_note.strip())
    finally:
        Path(tmp.name).unlink(missing_ok=True)


def _upload_result(results: List[dict], subject: str = "") -> HTMLResponse:
    row_parts = []
    for r in results:
        source_type = r.get("source_type", "")
        exam_hint = ""
        if source_type == "past_exam":
            scope_status = r.get("exam_scope_status", "unknown")
            usage_mode = r.get("exam_usage_mode", "style_generation")
            scope_label = EXAM_SCOPE_LABELS.get(scope_status, scope_status)
            usage_label = EXAM_USAGE_LABELS.get(usage_mode, usage_mode)
            exam_hint = (
                "<div class='hint'>"
                f"범위: {escape(scope_label)}"
                "<br/>"
                f"방식: {escape(usage_label)}"
                "</div>"
            )
        row_parts.append(
            "<tr>"
            f"<td><code>{escape(r.get('source_id', ''))}</code></td>"
            f"<td>{escape(r.get('title', ''))}</td>"
            f"<td>{escape(r.get('subject', ''))}</td>"
            f"<td><span class='pill'>{escape(source_type)}</span>{exam_hint}</td>"
            f"<td>{r.get('chunk_count', 0)}</td>"
            f"<td>{escape(r.get('extract_status', ''))}</td>"
            "</tr>"
        )
    rows = "".join(row_parts)
    return _page("업로드 완료", f"<section class='top'><div><h1>업로드 완료</h1><p>{len(results)}개 파일을 저장했습니다.</p></div><div class='nav'><a class='btn secondary' href='/upload'>추가 업로드</a><a class='btn' href='/sources/manage?subject={escape(subject)}'>파일 관리</a></div></section><section class='card'><table><thead><tr><th>source_id</th><th>제목</th><th>과목</th><th>유형</th><th>chunks</th><th>상태</th></tr></thead><tbody>{rows}</tbody></table></section>")



def _safe_download_headers(filename: str, mime_suffix: str) -> dict:
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", filename or "download").strip("._") or "download"
    if not base.lower().endswith(mime_suffix.lower()):
        base += mime_suffix
    utf8_name = quote((filename or "download") + ("" if (filename or "").lower().endswith(mime_suffix.lower()) else mime_suffix))
    return {"Content-Disposition": f'attachment; filename="{base}"; filename*=UTF-8\'\'{utf8_name}'}



def _md_is_table_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|.*\|\s*$", line or ""))


def _md_is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line or ""))


def _md_split_table_row(line: str) -> List[str]:
    raw = (line or "").strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [cell.strip() for cell in raw.split("|")]


def _inline_markdown_html(value: str) -> str:
    html = escape(value or "")
    html = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", html)
    html = re.sub(r"`([^`]+)`", r"<code>\1</code>", html)
    html = html.replace("&lt;mark&gt;", "<mark>").replace("&lt;/mark&gt;", "</mark>")
    return html


def _markdown_table_html(lines: List[str]) -> str:
    if len(lines) < 2:
        return ""
    header = _md_split_table_row(lines[0])
    body = [_md_split_table_row(line) for line in lines[2:] if _md_is_table_row(line)]
    col_count = max([len(header)] + [len(row) for row in body] + [1])
    def pad(row: List[str]) -> List[str]:
        return row + [""] * (col_count - len(row))
    head = "".join(f"<th>{_inline_markdown_html(cell)}</th>" for cell in pad(header))
    rows = []
    for row in body:
        rows.append("<tr>" + "".join(f"<td>{_inline_markdown_html(cell)}</td>" for cell in pad(row)) + "</tr>")
    return "<div class='table-wrap'><table><thead><tr>" + head + "</tr></thead><tbody>" + "".join(rows) + "</tbody></table></div>"


def _code_block_html(code: str, lang: str = "") -> str:
    visible = re.sub(r"[\s\u00a0\u200b\u200c\u200d\ufeff]+", " ", code or "").strip()
    if not visible or re.match(r"^[─━_\-=|+•·.\[\](){}\\/]+$", visible):
        return ""
    lower_lang = (lang or "").strip().lower()
    real_langs = {"python","py","javascript","js","typescript","ts","json","yaml","yml","html","css","sql","bash","sh","xml","toml","ini","prgm","casio"}
    looks_real = lower_lang in real_langs or bool(re.search(r"\b(function|class|def|import|from|const|let|var|return|SELECT|INSERT|UPDATE|DELETE|CREATE|for\s*\(|while\s*\()\b", code or ""))
    if not looks_real:
        return "<div class='flow-box'>" + "".join(f"<div>{_inline_markdown_html(line)}</div>" for line in (code or "").splitlines() if line.strip()) + "</div>"
    return f"<pre class='code-block'><code>{escape(code or '')}</code></pre>"


def _simple_markdown_html(markdown: str) -> str:
    lines = (markdown or "").replace("\r\n", "\n").split("\n")
    blocks: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            lang = line.replace("```", "", 1).strip()
            code: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            block = _code_block_html("\n".join(code), lang)
            if block:
                blocks.append(block)
            continue
        img = re.match(r"^!\[([^\]]*)\]\((.+)\)$", line.strip())
        if img:
            blocks.append(f"<figure class='image-card'><img alt='{escape(img.group(1) or '이미지')}' src='{escape(img.group(2))}'/><figcaption>{escape(img.group(1) or '이미지')}</figcaption></figure>")
            i += 1
            continue
        if _md_is_table_row(line) and i + 1 < len(lines) and _md_is_table_separator(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and _md_is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            blocks.append(_markdown_table_html(table_lines))
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            blocks.append(f"<h{level}>{_inline_markdown_html(heading.group(2))}</h{level}>")
            i += 1
            continue
        if re.match(r"^[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append("<li>" + _inline_markdown_html(re.sub(r"^[-*]\s+", "", lines[i])) + "</li>")
                i += 1
            blocks.append("<ul>" + "".join(items) + "</ul>")
            continue
        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6})\s+", lines[i]) and not re.match(r"^[-*]\s+", lines[i]) and not lines[i].startswith("```") and not re.match(r"^!\[", lines[i].strip()) and not (_md_is_table_row(lines[i]) and i + 1 < len(lines) and _md_is_table_separator(lines[i + 1])):
            para.append(lines[i])
            i += 1
        blocks.append("<p>" + _inline_markdown_html("\n".join(para)).replace("\n", "<br/>") + "</p>")
    return "\n".join(blocks)


def _clean_markdown_inline_for_docx(value: str) -> str:
    value = re.sub(r"`([^`]+)`", r"\1", value or "")
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = value.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    value = re.sub(r"<[^>]+>", "", value)
    return value


def _add_markdown_runs_to_docx(paragraph, line: str) -> None:
    from docx.enum.text import WD_COLOR_INDEX
    pos = 0
    for m in re.finditer(r"<mark>(.*?)</mark>", line or "", flags=re.S):
        if m.start() > pos:
            paragraph.add_run(_clean_markdown_inline_for_docx(line[pos:m.start()]))
        run = paragraph.add_run(_clean_markdown_inline_for_docx(m.group(1)))
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        pos = m.end()
    if pos < len(line or ""):
        paragraph.add_run(_clean_markdown_inline_for_docx((line or "")[pos:]))


def _add_markdown_table_to_docx(doc, table_lines: List[str]) -> None:
    if len(table_lines) < 2:
        return
    rows = [_md_split_table_row(table_lines[0])] + [_md_split_table_row(line) for line in table_lines[2:] if _md_is_table_row(line)]
    if not rows:
        return
    col_count = max(max(len(row) for row in rows), 1)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_markdown_runs_to_docx(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph("")


def _markdown_to_docx_bytes(title: str, markdown: str) -> BytesIO:
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.core_properties.title = title or "Study Note"
    doc.add_heading(title or "Study Note", level=1)
    image_pattern = re.compile(r"!\[([^\]]*)\]\((data:image/[^;]+;base64,([^)]+))\)")
    lines = (markdown or "").replace("\r\n", "\n").splitlines()
    fenced = False
    code_buffer: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip().startswith("```"):
            if fenced:
                code = "\n".join(code_buffer).strip()
                if code:
                    doc.add_paragraph(code, style="Intense Quote")
                code_buffer = []
                fenced = False
            else:
                fenced = True
            i += 1
            continue
        if fenced:
            code_buffer.append(line)
            i += 1
            continue
        if not line:
            doc.add_paragraph("")
            i += 1
            continue
        img = image_pattern.search(line)
        if img:
            try:
                bio = BytesIO(base64.b64decode(img.group(3), validate=False))
                doc.add_picture(bio, width=Inches(5.8))
                caption = _clean_markdown_inline_for_docx(img.group(1) or "이미지")
                if caption:
                    doc.add_paragraph(caption)
            except Exception:
                doc.add_paragraph("[이미지 삽입 실패: Study Note 화면에서 PDF 저장을 사용하세요]")
            i += 1
            continue
        if _md_is_table_row(line) and i + 1 < len(lines) and _md_is_table_separator(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and _md_is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table_to_docx(doc, table_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(_clean_markdown_inline_for_docx(heading.group(2)), level=level)
            i += 1
            continue
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs_to_docx(p, re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue
        p = doc.add_paragraph()
        _add_markdown_runs_to_docx(p, line)
        i += 1
    if fenced and code_buffer:
        code = "\n".join(code_buffer).strip()
        if code:
            doc.add_paragraph(code, style="Intense Quote")
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

@app.get("/health")
def health():
    return {"ok": True, "service": "lecturenote-suite", "version": "2.2.8"}


@app.get("/", response_class=HTMLResponse)
def root():
    return _page("LectureNote Suite", """
<section class="top"><div><h1>LectureNote Suite</h1><p>강의자료, 교재, 전사본, 시험지, 외부 정리본을 한 곳에 올리고 GPT Actions와 연결합니다.</p><div class="nav"><a class="btn" href="/upload">자료 업로드</a><a class="btn secondary" href="/static/study/index.html">정리본 보기</a><a class="btn secondary" href="/sources/manage">파일 관리</a><a class="btn secondary" href="/status">상태판</a><a class="btn secondary" href="/static/solvepad/index.html">SolvePad</a><a class="btn secondary" href="/static/casio/index.html">계산기 PRGM</a></div></div></section>
<section class="grid"><div class="card"><h2>통합 자료 업로드</h2><p>강의자료, 교재, 전사본, 시험지, GPT 생성 정리본, 외부 정리본을 자료 유형으로 구분해 한 화면에서 올립니다. 제목은 파일명으로 자동 지정됩니다.</p><div class="actions"><a class="btn" href="/upload">열기</a></div></div><div class="card"><h2>정리본 / Study Note</h2><p>GPT 생성 정리본, 외부 정리본, 시험 직전 정리, 계산기 사용법을 열고 수정·저장·내보내기합니다.</p><div class="actions"><a class="btn" href="/static/study/index.html">열기</a></div></div><div class="card"><h2>SolvePad 문제풀이</h2><p>GPT가 만든 문제팩을 iPad에서 불러오고 필기 풀이를 저장합니다.</p><div class="actions"><a class="btn" href="/static/solvepad/index.html">열기</a></div></div><div class="card"><h2>CASIO 계산기 PRGM</h2><p>GPT가 생성한 계산기 코드와 사용법/구조 해설을 확인합니다.</p><div class="actions"><a class="btn" href="/static/casio/index.html">열기</a></div></div><div class="card"><h2>관리 / 상태</h2><p>업로드 파일 삭제, 매핑 현황, API 문서를 확인합니다.</p><div class="actions"><a class="btn secondary" href="/sources/manage">파일 관리</a><a class="btn secondary" href="/status">상태판</a><a class="btn secondary" href="/status">매핑</a><a class="btn secondary" href="/docs">API</a></div></div></section>""")



@app.get("/status", response_class=HTMLResponse)
def status_page():
    return _page("LectureNote Suite 상태판", """
<section class="top"><div><h1>상태판 / 매핑 현황</h1><p>보호된 API를 직접 열지 않고, 저장된 액션 키로 자료 상태와 매핑 현황을 확인합니다.</p></div><div class="nav"><a class="btn secondary" href="/">홈</a><a class="btn secondary" href="/sources/manage">파일 관리</a><a class="btn secondary" href="/static/study/index.html">정리본 보기</a></div></section>
<section class="card"><label>액션 키</label><input name="action_key" type="password" placeholder="처음 한 번만 입력하면 이 브라우저에 자동 저장" autocomplete="off"/><div class="keybox"><span class="key-status" data-key-status></span><button class="btn secondary" type="button" data-clear-key>저장된 키 지우기</button></div><label>과목명</label><input name="subject" placeholder="예: CRE"/><div class="actions"><button type="button" id="loadStatus">상태 불러오기</button><a class="btn secondary" href="/docs">API 문서</a></div><div class="hint">홈 화면의 상태판/매핑 버튼은 이 화면을 통해 확인합니다. 직접 /dashboard 또는 /mapping/status를 열면 브라우저가 인증 헤더를 보낼 수 없어 Invalid action key가 날 수 있습니다.</div></section>
<section class="grid" style="margin-top:16px"><div class="card"><h2>요약</h2><div id="dashboardOut" class="muted">상태를 불러오세요.</div></div><div class="card"><h2>매핑 현황</h2><div id="mappingOut" class="muted">매핑 현황을 불러오세요.</div></div></section>
<script>
(function(){
  function el(id){return document.getElementById(id);}
  function key(){return (document.querySelector("input[name='action_key']")?.value || "").trim();}
  function subject(){return (document.querySelector("input[name='subject']")?.value || "").trim();}
  function esc(v){return String(v ?? '').replace(/[&<>"']/g, m => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[m]));}
  function headers(){return {Authorization: 'Bearer ' + key()};}
  function renderJsonBox(obj){
    return '<pre style="white-space:pre-wrap;overflow:auto;max-height:480px;background:#f8fafc;border:1px solid #e5e7eb;border-radius:14px;padding:12px">' + esc(JSON.stringify(obj, null, 2)) + '</pre>';
  }
  function dashHtml(data){
    const sources = data.sources || data.source_summary || data.sourcesByType || {};
    const notes = data.study_notes || data.notes || [];
    const runs = data.workflow_runs || data.runs || [];
    const calc = data.calculator_projects || [];
    let html = '<ul>';
    html += '<li>정리본: ' + esc(notes.length ?? 0) + '개</li>';
    html += '<li>계산기 프로젝트: ' + esc(calc.length ?? 0) + '개</li>';
    html += '<li>진행 작업: ' + esc(runs.length ?? 0) + '개</li>';
    html += '<li>자료 요약: ' + esc(typeof sources === 'object' ? Object.keys(sources).length + '개 항목' : sources) + '</li>';
    html += '</ul>';
    return html + '<details><summary>원본 JSON</summary>' + renderJsonBox(data) + '</details>';
  }
  function mappingHtml(data){
    const units = data.units || data.mapped_units || data.unit_maps || [];
    const unmapped = data.unmapped_sources || data.unmapped || [];
    let html = '<ul>';
    html += '<li>매핑 단원/맵: ' + esc(units.length ?? 0) + '개</li>';
    html += '<li>미매핑 자료: ' + esc(unmapped.length ?? 0) + '개</li>';
    html += '</ul>';
    return html + '<details><summary>원본 JSON</summary>' + renderJsonBox(data) + '</details>';
  }
  async function load(){
    const k = key();
    if(!k){
      alert('액션 키를 입력하세요. Render 환경변수 ACTION_API_KEY와 같은 값입니다.');
      return;
    }
    if(subject()) localStorage.setItem('lecturenote_subject', subject());
    localStorage.setItem('lecturenote_action_key', k);
    el('dashboardOut').textContent = '불러오는 중...';
    el('mappingOut').textContent = '불러오는 중...';
    const qs = subject() ? '?subject=' + encodeURIComponent(subject()) : '';
    try {
      const [dashRes, mapRes] = await Promise.all([
        fetch('/dashboard' + qs, {headers: headers()}),
        fetch('/mapping/status' + qs, {headers: headers()})
      ]);
      if(!dashRes.ok) throw new Error(await dashRes.text());
      if(!mapRes.ok) throw new Error(await mapRes.text());
      const dash = await dashRes.json();
      const map = await mapRes.json();
      el('dashboardOut').innerHTML = dashHtml(dash);
      el('mappingOut').innerHTML = mappingHtml(map);
    } catch(e) {
      el('dashboardOut').textContent = '오류';
      el('mappingOut').textContent = '오류';
      alert(String(e.message || e));
    }
  }
  document.addEventListener('DOMContentLoaded', function(){
    el('loadStatus')?.addEventListener('click', load);
  });
})();
</script>""")


@app.get("/upload", response_class=HTMLResponse)
def upload_page():
    options = _options('lecture_slides')
    body = """
<section class="top"><div><h1>통합 자료 업로드</h1><p>강의자료, 교재, 전사본, 시험지, 외부 정리본을 자료 유형으로 골라 한 번에 업로드합니다.</p></div><div class="nav"><a class="btn secondary" href="/">홈</a><a class="btn secondary" href="/sources/manage">파일 관리</a></div></section>
<section class="card"><form action="/sources/upload-batch" method="post" enctype="multipart/form-data"><label>액션 키</label><input name="action_key" type="password" placeholder="처음 한 번만 입력하면 이 브라우저에 자동 저장" autocomplete="off"/><div class="keybox"><span class="key-status" data-key-status></span><button class="btn secondary" type="button" data-clear-key>저장된 키 지우기</button></div><div class="hint">보안 때문에 서버 전체 공개 업로드는 막아두고, 키는 이 브라우저 localStorage에만 저장합니다.</div><label>과목명</label><input name="subject" placeholder="예: CRE, 반응공학, 수학2"/><div class="hint">선택사항이지만 과목별 관리에는 입력 권장.</div><label>자료 유형</label><select name="source_type" id="sourceTypeSelect">__OPTIONS__</select><div id="examMetaBox" class="exam-meta" hidden><h2>시험지/기출 사용 설정</h2><p class="hint">기출 업로드 시 현재 시험범위 해당 여부와 나중에 문제를 그대로 전사할지, 출제 스타일만 분석해 새 문제를 만들지 저장합니다.</p><label>현재 시험범위 해당 여부</label><select name="exam_scope_status"><option value="unknown">확인 안 됨</option><option value="in_scope">현재 시험범위 해당</option><option value="out_of_scope">현재 시험범위 아님</option><option value="mixed">일부만 해당</option></select><label>문제 사용 방식</label><select name="exam_usage_mode"><option value="style_generation">출제 스타일 분석 후 새 문제 생성</option><option value="exact_transcription">해당 시험지 문제 그대로 전사</option><option value="both">그대로 전사 + 새 문제 생성 둘 다</option><option value="exclude">참고 제외/보관만</option></select><label>범위/사용 메모</label><input name="exam_range_note" placeholder="예: 1~3단원만 해당, 2025 기말 형식 참고, 계산형 문제만 전사 등"/></div><label>제목</label><input name="title" placeholder="선택사항. 비우면 각 파일명으로 자동 저장"/><label>파일</label><input type="file" name="files" multiple required/><div class="hint">여러 파일을 한 번에 선택 가능합니다.</div><div class="actions"><button type="submit">업로드</button><a class="btn secondary" href="/">홈으로</a></div></form></section><script>document.addEventListener('DOMContentLoaded',function(){const sel=document.getElementById('sourceTypeSelect');const box=document.getElementById('examMetaBox');function sync(){if(box) box.hidden = sel.value !== 'past_exam';} if(sel){sel.addEventListener('change',sync); sync();}});</script>
""".replace("__OPTIONS__", options)
    return _page("자료 업로드", body)


@app.get("/notes/upload", response_class=HTMLResponse)
def notes_upload_page():
    return _page("외부 정리본 업로드", """
<section class="top"><div><h1>외부 정리본 업로드</h1><p>이제 외부 정리본도 통합 자료 업로드 화면의 자료 유형에서 선택합니다.</p></div><div class="nav"><a class="btn" href="/upload">통합 업로드로 이동</a><a class="btn secondary" href="/">홈</a></div></section>
<section class="card"><p>업로드 화면에서 <b>자료 유형</b>을 <code>외부 정리본</code>으로 선택하면 됩니다. 기존 링크 호환을 위해 이 페이지는 남겨둡니다.</p></section>""")


@app.post("/sources/upload")
async def upload_source(source_type: str = Form(default="lecture_slides"), title: str = Form(default=""), subject: str = Form(default=""), action_key: str = Form(default=""), exam_scope_status: str = Form(default="unknown"), exam_usage_mode: str = Form(default="style_generation"), exam_range_note: str = Form(default=""), file: UploadFile = File(...), authorization: Optional[str] = Header(default=None), x_action_key: Optional[str] = Header(default=None)):
    if not _is_authorized(authorization, x_action_key, action_key):
        raise HTTPException(status_code=401, detail="Invalid action key")
    return _upload_result([await _save_uploaded_file(file, source_type, subject, title, exam_scope_status, exam_usage_mode, exam_range_note)], subject.strip())


@app.post("/sources/upload-batch")
async def upload_sources_batch(source_type: str = Form(default="lecture_slides"), title: str = Form(default=""), subject: str = Form(default=""), action_key: str = Form(default=""), exam_scope_status: str = Form(default="unknown"), exam_usage_mode: str = Form(default="style_generation"), exam_range_note: str = Form(default=""), files: List[UploadFile] = File(...), authorization: Optional[str] = Header(default=None), x_action_key: Optional[str] = Header(default=None)):
    if not _is_authorized(authorization, x_action_key, action_key):
        raise HTTPException(status_code=401, detail="Invalid action key")
    results = []
    for i, file in enumerate(files, 1):
        item_title = f"{title.strip()} {i}" if title.strip() and len(files) > 1 else title.strip()
        results.append(await _save_uploaded_file(file, source_type, subject, item_title, exam_scope_status, exam_usage_mode, exam_range_note))
    return _upload_result(results, subject.strip())


@app.post("/sources/text", dependencies=[Depends(require_auth)])
def save_text_source_endpoint(payload: TextSourceSave):
    return save_text_source(payload.title, payload.source_type, payload.text, payload.original_name, subject=payload.subject.strip())


@app.post("/notes/text", dependencies=[Depends(require_auth)])
def save_external_note_endpoint(payload: TextSourceSave):
    source_type = payload.source_type if payload.source_type in {"external_note", "generated_note", "note_example"} else "external_note"
    return save_text_source(payload.title, source_type, payload.text, payload.original_name, subject=payload.subject.strip())


@app.get("/sources")
def sources(source_type: str = Query(default=""), subject: str = Query(default="")):
    return {"sources": list_sources(source_type, subject)}


@app.get("/sources/unmapped", dependencies=[Depends(require_auth)])
def unmapped_sources(source_type: str = Query(default=""), subject: str = Query(default="")):
    return {"unmapped_sources": list_unmapped_sources([source_type] if source_type else [], subject)}


@app.get("/external-notes")
def external_notes_endpoint(subject: str = Query(default="")):
    return {"external_notes": list_external_notes(subject)}


@app.get("/sources/manage", response_class=HTMLResponse)
def manage_sources_page(subject: str = Query(default=""), action_key: str = Query(default="")):
    rows = []
    for s in list_sources(subject=subject):
        raw_sid = s.get("id", "")
        sid = escape(raw_sid)
        rows.append(
            "<tr>"
            f"<td><input type='checkbox' name='source_ids' value='{sid}' data-source-check aria-label='자료 선택'/></td>"
            f"<td><span class='pill'>{escape(s.get('subject','') or '미지정')}</span></td>"
            f"<td>{escape(s.get('source_type',''))}</td>"
            f"<td>{escape(s.get('title',''))}<div class='hint'><code>{sid}</code></div></td>"
            f"<td>{escape(s.get('original_name',''))}</td>"
            f"<td>{escape(s.get('created_at',''))}</td>"
            f"<td><button class='danger' type='submit' name='source_ids' value='{sid}' onclick=\"return confirm('이 자료를 삭제할까요?');\">삭제</button></td>"
            "</tr>"
        )
    table_rows = "\n".join(rows) or "<tr><td colspan='7' class='muted'>업로드된 자료가 없습니다.</td></tr>"
    return _page(
        "업로드 파일 관리",
        f"""<section class='top'><div><h1>업로드 파일 관리</h1><p>업로드한 자료를 과목별로 확인하고 삭제할 수 있습니다.</p></div><div class='nav'><a class='btn secondary' href='/upload'>자료 업로드</a><a class='btn secondary' href='/'>홈</a></div></section>
<section class='card'><form method='get' action='/sources/manage'><label>과목 필터</label><input name='subject' value='{escape(subject)}' placeholder='예: CRE'/><label>액션 키</label><input name='action_key' type='password' value='{escape(action_key)}' placeholder='처음 한 번 입력하면 자동 저장'/><div class='keybox'><span class='key-status' data-key-status></span><button class='btn secondary' type='button' data-clear-key>저장된 키 지우기</button></div><div class='actions'><button type='submit'>적용</button><a class='btn secondary' href='/status?subject={escape(subject)}'>상태판/매핑</a></div></form></section>
<section class='card' style='margin-top:16px;overflow:auto'><form action='/sources/delete-batch' method='post' onsubmit="const n=document.querySelectorAll('[data-source-check]:checked').length; if(!n && !event.submitter?.value){{alert('삭제할 자료를 선택하세요.'); return false;}} return confirm((n || 1)+'개 자료를 삭제할까요?');"><input type='hidden' name='action_key' value='{escape(action_key)}'/><input type='hidden' name='subject' value='{escape(subject)}'/><div class='actions' style='justify-content:space-between;align-items:center;margin-top:0;margin-bottom:12px'><label style='margin:0;display:flex;gap:8px;align-items:center'><input type='checkbox' id='selectAllSources' onclick="document.querySelectorAll('[data-source-check]').forEach(cb=>cb.checked=this.checked)"/> 전체 선택</label><button class='danger' type='submit'>선택 삭제</button></div><table><thead><tr><th>선택</th><th>과목</th><th>유형</th><th>제목/source_id</th><th>파일</th><th>생성일</th><th>작업</th></tr></thead><tbody>{table_rows}</tbody></table></form></section>""",
    )


@app.get("/sources/{source_id}")
def source_detail(source_id: str):
    source = get_source(source_id)
    if not source:
        raise HTTPException(status_code=404, detail="Source not found")
    return source


@app.get("/dashboard", dependencies=[Depends(require_auth)])
def dashboard_endpoint(subject: str = Query(default=""), limit: int = Query(default=8)):
    return get_dashboard(subject, limit)


@app.get("/mapping/status", dependencies=[Depends(require_auth)])
def mapping_status_endpoint(subject: str = Query(default="")):
    return get_mapping_status(subject)


@app.delete("/sources/{source_id}", dependencies=[Depends(require_auth)])
def delete_source_endpoint(source_id: str):
    result = delete_source(source_id)
    if not result:
        raise HTTPException(status_code=404, detail="Source not found")
    return result


@app.post("/sources/{source_id}/delete")
def delete_source_form_endpoint(source_id: str, action_key: str = Form(default="")):
    if not _is_authorized(action_key=action_key):
        raise HTTPException(status_code=401, detail="Invalid action key")
    result = delete_source(source_id)
    if not result:
        raise HTTPException(status_code=404, detail="Source not found")
    return _page("삭제 완료", f"<section class='top'><div><h1>삭제 완료</h1><p>자료와 추출 chunk를 삭제했습니다.</p></div><div class='nav'><a class='btn' href='/sources/manage?action_key={escape(action_key)}'>파일 관리로 돌아가기</a></div></section><section class='card'><p><b>제목:</b> {escape(result.get('deleted_title',''))}</p><p><b>source_id:</b> <code>{escape(result.get('deleted_source_id',''))}</code></p><p><b>삭제 chunks:</b> {result.get('deleted_chunks')}</p><p><b>파일 삭제:</b> {result.get('deleted_file')}</p></section>")


@app.post("/sources/delete-batch")
def delete_sources_batch_endpoint(action_key: str = Form(default=""), subject: str = Form(default=""), source_ids: Optional[List[str]] = Form(default=None)):
    if not _is_authorized(action_key=action_key):
        raise HTTPException(status_code=401, detail="Invalid action key")
    ids = []
    seen = set()
    for sid in source_ids or []:
        clean = (sid or "").strip()
        if clean and clean not in seen:
            seen.add(clean)
            ids.append(clean)
    if not ids:
        return _page("삭제할 자료 없음", f"<section class='top'><div><h1>삭제할 자료가 없습니다</h1><p>파일 관리에서 삭제할 자료를 먼저 선택하세요.</p></div><div class='nav'><a class='btn' href='/sources/manage?action_key={escape(action_key)}&subject={escape(subject)}'>파일 관리로 돌아가기</a></div></section>")
    deleted = []
    missing = []
    for sid in ids:
        result = delete_source(sid)
        if result:
            deleted.append(result)
        else:
            missing.append(sid)
    rows = "".join(
        "<tr>"
        f"<td>{escape(item.get('deleted_title',''))}</td>"
        f"<td><code>{escape(item.get('deleted_source_id',''))}</code></td>"
        f"<td>{item.get('deleted_chunks', 0)}</td>"
        f"<td>{escape(str(item.get('deleted_file', '')))}</td>"
        "</tr>"
        for item in deleted
    )
    if not rows:
        rows = "<tr><td colspan='4' class='muted'>삭제된 자료가 없습니다.</td></tr>"
    missing_html = ""
    if missing:
        missing_html = "<p class='hint'>찾을 수 없는 source_id: " + ", ".join(f"<code>{escape(x)}</code>" for x in missing) + "</p>"
    return _page(
        "선택 삭제 완료",
        f"<section class='top'><div><h1>선택 삭제 완료</h1><p>{len(deleted)}개 자료를 삭제했습니다.</p>{missing_html}</div><div class='nav'><a class='btn' href='/sources/manage?action_key={escape(action_key)}&subject={escape(subject)}'>파일 관리로 돌아가기</a><a class='btn secondary' href='/'>홈</a></div></section><section class='card'><table><thead><tr><th>제목</th><th>source_id</th><th>삭제 chunks</th><th>파일 삭제</th></tr></thead><tbody>{rows}</tbody></table></section>",
    )


@app.post("/sources/search", dependencies=[Depends(require_auth)])
def search(payload: SearchRequest):
    return {"results": search_sources(payload.query, payload.source_types, payload.limit, payload.subject)}


@app.get("/workflow/options", dependencies=[Depends(require_auth)])
def workflow_options(subject: str = Query(default="")):
    return get_workflow_options(subject)


@app.post("/workflow/plans", dependencies=[Depends(require_auth)])
def create_workflow_plan_endpoint(payload: WorkflowPlanCreate):
    return create_workflow_plan(payload.title, payload.subject, payload.selected_units, payload.selected_mode, payload.unit_map_id, payload.source_ids, payload.reference_priority, payload.notes)


@app.get("/workflow/plans", dependencies=[Depends(require_auth)])
def list_workflow_plans_endpoint():
    return {"workflow_plans": list_workflow_plans()}


@app.post("/workflow/runs", dependencies=[Depends(require_auth)])
def create_workflow_run_endpoint(payload: WorkflowRunCreate):
    return create_workflow_run(payload.title, payload.mode, payload.subject, payload.selected_units, payload.workflow_plan_id, payload.total_steps, payload.metadata)


@app.get("/workflow/runs", dependencies=[Depends(require_auth)])
def list_workflow_runs_endpoint(status: str = Query(default="")):
    return {"workflow_runs": list_workflow_runs(status)}


@app.get("/workflow/runs/{run_id}/next", dependencies=[Depends(require_auth)])
def get_next_workflow_step_endpoint(run_id: str):
    result = get_next_workflow_step(run_id)
    if not result:
        raise HTTPException(status_code=404, detail="Workflow run not found")
    return result


@app.post("/workflow/checkpoints", dependencies=[Depends(require_auth)])
def save_workflow_checkpoint_endpoint(payload: WorkflowCheckpointSave):
    result = save_workflow_checkpoint(payload.run_id, payload.step_index, payload.step_label, payload.status, payload.saved_refs, payload.next_action, payload.notes, payload.advance_to_next)
    if not result:
        raise HTTPException(status_code=404, detail="Workflow run not found")
    return result


@app.get("/workflow/runs/{run_id}/checkpoints", dependencies=[Depends(require_auth)])
def list_workflow_checkpoints_endpoint(run_id: str):
    return {"checkpoints": list_workflow_checkpoints(run_id)}


@app.post("/projects", dependencies=[Depends(require_auth)])
def create_project_endpoint(payload: ProjectCreate):
    return create_project(payload.title, payload.project_type, payload.source_ids, payload.metadata)


@app.post("/projects/{project_id}/outline", dependencies=[Depends(require_auth)])
def save_outline_endpoint(project_id: str, payload: OutlineSave):
    return save_outline(project_id, [section.model_dump() for section in payload.sections])


@app.get("/projects/{project_id}/next", dependencies=[Depends(require_auth)])
def next_section(project_id: str):
    return get_next_section(project_id)


@app.post("/projects/{project_id}/sections", dependencies=[Depends(require_auth)])
def save_section_endpoint(project_id: str, payload: SectionSave):
    return save_section(project_id, payload.model_dump())


@app.post("/projects/{project_id}/items/{item_type}", dependencies=[Depends(require_auth)])
def save_items_endpoint(project_id: str, item_type: str, payload: GenericItemsSave):
    return save_project_items(project_id, item_type, payload.items)


@app.get("/projects/{project_id}/final-bundle", dependencies=[Depends(require_auth)])
def final_bundle_endpoint(project_id: str):
    bundle = final_bundle(project_id)
    if not bundle:
        raise HTTPException(status_code=404, detail="Project not found")
    return bundle


@app.post("/projects/{project_id}/export-note-source", dependencies=[Depends(require_auth)])
def export_project_note_source_endpoint(project_id: str, payload: dict = None):
    result = export_project_note_as_source(project_id, (payload or {}).get("title", ""))
    if not result:
        raise HTTPException(status_code=404, detail="Project not found")
    return result


@app.post("/notes/versions", dependencies=[Depends(require_auth)])
def save_note_version_endpoint(payload: NoteVersionSave):
    return save_note_version(payload.title, payload.content_markdown, payload.series_id, payload.source_type, payload.change_summary, payload.based_on_version, payload.subject.strip(), payload.replace_latest)


@app.get("/notes/versions", dependencies=[Depends(require_auth)])
def list_note_versions_endpoint(series_id: str = Query(default="")):
    return {"versions": list_note_versions(series_id)}


@app.get("/notes/versions/{series_id}/latest", dependencies=[Depends(require_auth)])
def latest_note_version_endpoint(series_id: str):
    latest = get_latest_note_version(series_id)
    if not latest:
        raise HTTPException(status_code=404, detail="Note series not found")
    return latest


@app.post("/transcripts/revisions", dependencies=[Depends(require_auth)])
def save_transcript_revision_endpoint(payload: TranscriptRevisionSave):
    return save_transcript_revision(payload.title, payload.corrected_text, payload.original_transcript_source_id, payload.terminology_map, payload.change_log, payload.subject.strip())


@app.post("/unit-maps", dependencies=[Depends(require_auth)])
def save_unit_map_endpoint(payload: dict):
    title = str(payload.get("title") or "Unit Map")
    source_ids = payload.get("source_ids") or payload.get("sourceIds") or []
    if isinstance(source_ids, str):
        source_ids = [source_ids]
    if not isinstance(source_ids, list):
        source_ids = []
    map_json = payload.get("map") or payload.get("map_json") or payload.get("mapJson") or payload.get("mapping")
    if not isinstance(map_json, dict):
        raise HTTPException(status_code=422, detail="Unit map payload must include object field: map")
    created_by = str(payload.get("created_by") or payload.get("createdBy") or "gpt")
    return save_unit_map(title, source_ids, map_json, created_by)


@app.get("/unit-maps")
def list_unit_maps_endpoint():
    return {"unit_maps": list_unit_maps()}



@app.get("/study/notes", dependencies=[Depends(require_auth)])
def list_study_notes_endpoint(subject: str = Query(default=""), source_type: str = Query(default="")):
    return {"notes": list_study_notes(subject, source_type)}


@app.post("/study/notes", dependencies=[Depends(require_auth)])
def save_study_note_endpoint(payload: StudyNoteSave):
    source_type = payload.source_type if payload.source_type in {"generated_note", "external_note", "exam_cram", "calculator_manual", "note_example"} else "generated_note"
    return save_note_version(
        payload.title,
        payload.content_markdown,
        payload.series_id,
        source_type,
        payload.change_summary,
        None,
        payload.subject.strip(),
        payload.replace_latest,
    )


@app.post("/exam-cram", dependencies=[Depends(require_auth)])
def save_exam_cram_endpoint(payload: ExamCramSave):
    return save_note_version(
        payload.title,
        payload.content_markdown,
        payload.series_id,
        "exam_cram",
        payload.change_summary,
        None,
        payload.subject.strip(),
        payload.replace_latest,
    )


@app.get("/study/notes/{source_id}", dependencies=[Depends(require_auth)])
def get_study_note_endpoint(source_id: str):
    note = get_source_markdown(source_id)
    if not note:
        raise HTTPException(status_code=404, detail="Study note not found")
    return note


@app.put("/study/notes/{source_id}", dependencies=[Depends(require_auth)])
def update_study_note_endpoint(source_id: str, payload: StudyNoteUpdate):
    result = update_source_markdown(source_id, payload.content_markdown, payload.title, payload.change_summary)
    if not result:
        raise HTTPException(status_code=404, detail="Study note not found")
    return result



@app.get("/sources/{source_id}/slide-image", dependencies=[Depends(require_auth)])
def render_source_slide_image(source_id: str, page: int = Query(default=1, ge=1), zoom: float = Query(default=2.0, ge=0.5, le=4.0)):
    source = get_source(source_id)
    if not source:
        raise HTTPException(status_code=404, detail="Source not found")
    stored_path = Path(source.get("stored_path") or "")
    if not stored_path.exists():
        raise HTTPException(status_code=404, detail="Stored file not found")
    suffix = stored_path.suffix.lower()
    if suffix != ".pdf" and "pdf" not in (source.get("mime_type") or "").lower():
        raise HTTPException(status_code=400, detail="Slide image rendering currently supports PDF lecture slides only")
    try:
        import fitz  # PyMuPDF
    except Exception:
        raise HTTPException(status_code=500, detail="PDF slide rendering package is not installed. Add PyMuPDF to requirements.txt and redeploy.")
    try:
        doc = fitz.open(str(stored_path))
        if page > len(doc):
            raise HTTPException(status_code=400, detail=f"Page out of range. This PDF has {len(doc)} pages.")
        pdf_page = doc[page - 1]
        matrix = fitz.Matrix(zoom, zoom)
        pix = pdf_page.get_pixmap(matrix=matrix, alpha=False)
        png = pix.tobytes("png")
        data_url = "data:image/png;base64," + base64.b64encode(png).decode("ascii")
        label = f"{source.get('title') or source.get('original_name') or source_id} p.{page}"
        return {
            "source_id": source_id,
            "page": page,
            "page_count": len(doc),
            "label": label,
            "data_url": data_url,
            "markdown": f"![{label}]({data_url})",
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to render slide page: {exc}")

@app.get("/study/notes/{source_id}/download.md")
def download_study_note_md(source_id: str):
    note = get_source_markdown(source_id)
    if not note:
        raise HTTPException(status_code=404, detail="Study note not found")
    title = note["source"].get("title", source_id).replace('/', '_')
    return Response(
        note.get("markdown", ""),
        media_type="text/markdown; charset=utf-8",
        headers=_safe_download_headers(title, ".md"),
    )


@app.get("/study/notes/{source_id}/download.docx")
def download_study_note_docx(source_id: str):
    note = get_source_markdown(source_id)
    if not note:
        raise HTTPException(status_code=404, detail="Study note not found")
    title = note["source"].get("title", source_id).replace('/', '_')
    try:
        out = _markdown_to_docx_bytes(title, note.get("markdown", ""))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {exc}")
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=_safe_download_headers(title, ".docx"),
    )


@app.get("/study/notes/{source_id}/print", response_class=HTMLResponse)
def print_study_note_page(source_id: str):
    note = get_source_markdown(source_id)
    if not note:
        raise HTTPException(status_code=404, detail="Study note not found")
    title = escape(note["source"].get("title", source_id))
    html = _simple_markdown_html(note.get("markdown", ""))
    return HTMLResponse(f"""<!doctype html><html lang='ko'><head><meta charset='utf-8'/><meta name='viewport' content='width=device-width,initial-scale=1'/><title>{title}</title><script>window.MathJax={{tex:{{inlineMath:[["$","$"],["\\\\(","\\\\)"]],displayMath:[["$$","$$"],["\\\\[","\\\\]"]],processEscapes:true}},svg:{{fontCache:'global'}}}};</script><script defer src='https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js'></script><style>body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI','Noto Sans KR',Arial,sans-serif;line-height:1.72;max-width:820px;margin:0 auto;padding:34px;color:#111827}}.bar{{display:flex;gap:8px;margin-bottom:20px}}button,a{{border:0;border-radius:10px;padding:9px 12px;background:#4f46e5;color:white;text-decoration:none;font-weight:800}}img{{max-width:100%;border-radius:10px}}figure{{margin:18px 0;text-align:center}}figcaption{{font-size:12px;color:#667085;margin-top:6px}}mark{{background:#fff39a}}.table-wrap{{overflow:visible;margin:18px 0}}table{{width:100%;border-collapse:collapse;font-size:13px}}th,td{{border:1px solid #d0d5dd;padding:8px 10px;vertical-align:top;text-align:left}}th{{background:#f2f4f7;font-weight:800}}pre,.flow-box{{background:#f8fafc;border:1px solid #e4e7ec;border-radius:10px;padding:12px;white-space:pre-wrap}}@media print{{.bar{{display:none}}body{{padding:0;max-width:none}}table{{break-inside:auto}}tr{{break-inside:avoid}}}}</style></head><body><div class='bar'><button onclick='window.print()'>PDF로 저장/인쇄</button><a href='/static/study/index.html'>Study Studio</a><a href='/'>홈</a></div><article>{html}</article></body></html>""")

@app.get("/unit-maps/{unit_map_id}", dependencies=[Depends(require_auth)])
def get_unit_map_endpoint(unit_map_id: str):
    unit_map = get_unit_map(unit_map_id)
    if not unit_map:
        raise HTTPException(status_code=404, detail="Unit map not found")
    return unit_map


@app.post("/problem-packs", dependencies=[Depends(require_auth)])
def save_problem_pack_endpoint(payload: ProblemPackSave):
    saved = save_problem_pack(payload.title, payload.pack)
    return {**saved, "import_url": f"{PUBLIC_BASE_URL}/static/solvepad/index.html?server={PUBLIC_BASE_URL}&importToken={saved['token']}"}


@app.get("/packs/{token}")
def get_pack_for_solvepad(token: str):
    pack = get_problem_pack_by_token(token)
    if not pack:
        raise HTTPException(status_code=404, detail="Problem pack not found")
    return pack


@app.get("/calculator/projects", dependencies=[Depends(require_auth)])
def list_calculator_projects_endpoint(subject: str = Query(default="")):
    return {"calculator_projects": list_calculator_blueprints(subject)}


@app.get("/calculator/projects/{project_id}", dependencies=[Depends(require_auth)])
def get_calculator_project_endpoint(project_id: str):
    project = get_calculator_blueprint(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Calculator project not found")
    return project


@app.get("/calculator/projects/{project_id}/manual", response_class=HTMLResponse)
def get_calculator_manual_page(project_id: str):
    project = get_calculator_blueprint(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Calculator project not found")
    manual = escape(project.get("manual_markdown", "") or "사용법 문서가 저장되지 않았습니다.")
    title = escape(project.get("title", project_id))
    return _page(f"계산기 사용법 - {title}", f"<section class='top'><div><h1>{title}</h1><p>CASIO PRGM 사용법 및 구조 해설</p></div><div class='nav'><a class='btn secondary' href='/static/casio/index.html'>계산기 Studio</a><a class='btn secondary' href='/'>홈</a></div></section><section class='card'><pre style='white-space:pre-wrap;line-height:1.65'>{manual}</pre></section>")


@app.delete("/calculator/projects/{project_id}", dependencies=[Depends(require_auth)])
def delete_calculator_project_endpoint(project_id: str):
    result = delete_calculator_blueprint(project_id)
    if not result:
        raise HTTPException(status_code=404, detail="Calculator project not found")
    return result



@app.get("/calculator/projects/{project_id}/download.zip", dependencies=[Depends(require_auth)])
def download_calculator_project_zip(project_id: str):
    project = get_calculator_blueprint(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Calculator project not found")
    mem = BytesIO()
    with zipfile.ZipFile(mem, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file_info in (project.get("generated") or {}).get("files", []):
            archive.writestr(file_info.get("name", "PROGRAM.TXT"), file_info.get("content", ""))
        if project.get("manual_markdown"):
            archive.writestr("MANUAL.md", project.get("manual_markdown", ""))
        if project.get("analysis_markdown"):
            archive.writestr("ANALYSIS.md", project.get("analysis_markdown", ""))
    mem.seek(0)
    return StreamingResponse(mem, media_type="application/zip", headers={"Content-Disposition": f'attachment; filename="{project_id}-casio.zip"'})

@app.post("/calculator/validate", dependencies=[Depends(require_auth)])
def validate_calculator(payload: CalculatorBlueprintSave):
    return validate_blueprint(payload.blueprint).as_dict()


@app.post("/calculator/generate", dependencies=[Depends(require_auth)])
def generate_calculator(payload: CalculatorBlueprintSave):
    validation = validate_blueprint(payload.blueprint).as_dict()
    generated = generate_txt_files(payload.blueprint).as_dict()
    saved = save_calculator_blueprint(
        payload.title,
        payload.blueprint,
        validation,
        generated,
        payload.metadata,
        payload.manual_markdown,
        payload.analysis_markdown,
        payload.replace_calculator_project_id,
    )
    project_id = saved["calculator_project_id"]
    return {
        **saved,
        "validation": validation,
        "generated": generated,
        "manual_url": f"{PUBLIC_BASE_URL}/calculator/projects/{project_id}/manual",
        "studio_url": f"{PUBLIC_BASE_URL}/static/casio/index.html?projectId={project_id}",
    }


@app.post("/calculator/generate.zip", dependencies=[Depends(require_auth)])
def generate_calculator_zip(payload: CalculatorBlueprintSave):
    generated = generate_txt_files(payload.blueprint).as_dict()
    mem = BytesIO()
    with zipfile.ZipFile(mem, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file_info in generated.get("files", []):
            archive.writestr(file_info["name"], file_info["content"])
        if payload.manual_markdown.strip():
            archive.writestr("MANUAL.md", payload.manual_markdown)
        if payload.analysis_markdown.strip():
            archive.writestr("ANALYSIS.md", payload.analysis_markdown)
    mem.seek(0)
    return StreamingResponse(mem, media_type="application/zip", headers={"Content-Disposition": 'attachment; filename="casio-prgm-files.zip"'})
