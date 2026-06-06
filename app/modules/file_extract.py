from __future__ import annotations

from pathlib import Path
from typing import Tuple

TEXT_EXTS = {
    ".txt", ".md", ".csv", ".tsv", ".json", ".yaml", ".yml",
    ".py", ".js", ".html", ".css", ".tex", ".log"
}

MAX_EXTRACT_CHARS = 500_000


def _clip(text: str) -> str:
    if len(text) > MAX_EXTRACT_CHARS:
        return text[:MAX_EXTRACT_CHARS] + "\n\n[TRUNCATED: extracted text exceeded server limit]"
    return text


def _read_text(path: Path) -> Tuple[str, str]:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp949", "latin-1"):
        try:
            return _clip(raw.decode(enc)), "ok"
        except UnicodeDecodeError:
            continue
    return "", "decode_failed"


def extract_text(path: Path, filename: str, mime_type: str = "") -> Tuple[str, str]:
    """Return (text, status). Keeps extraction local and deterministic.

    status values:
    - ok
    - ok_empty
    - unsupported_binary
    - pdf_extracted / pdf_extract_failed / pdf_package_missing
    - docx_extracted / docx_extract_failed / docx_package_missing
    - decode_failed
    """
    suffix = Path(filename or path.name).suffix.lower()
    mime = (mime_type or "").lower()

    if suffix in TEXT_EXTS or mime.startswith("text/") or "json" in mime or "csv" in mime:
        text, status = _read_text(path)
        return text, status if text else "ok_empty"

    if suffix == ".pdf" or mime == "application/pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            return "", "pdf_package_missing"
        try:
            reader = PdfReader(str(path))
            pages = []
            for i, page in enumerate(reader.pages):
                try:
                    pages.append(f"\n\n--- Page {i+1} ---\n" + (page.extract_text() or ""))
                except Exception:
                    pages.append(f"\n\n--- Page {i+1} ---\n[PAGE TEXT EXTRACTION FAILED]")
            text = _clip("".join(pages).strip())
            return text, "pdf_extracted" if text else "ok_empty"
        except Exception as exc:
            return f"[PDF extraction failed: {exc}]", "pdf_extract_failed"

    if suffix == ".docx" or "wordprocessingml" in mime:
        try:
            from docx import Document  # type: ignore
        except Exception:
            return "", "docx_package_missing"
        try:
            doc = Document(str(path))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    parts.append(" | ".join(cells))
            text = _clip("\n".join(parts).strip())
            return text, "docx_extracted" if text else "ok_empty"
        except Exception as exc:
            return f"[DOCX extraction failed: {exc}]", "docx_extract_failed"

    return "", "unsupported_binary"
